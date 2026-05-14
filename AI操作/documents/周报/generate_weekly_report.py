#!/usr/bin/env python3
# -*- coding: utf-8 -*-

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from datetime import datetime

def create_weekly_report():
    doc = Document()

    # 设置标题
    title = doc.add_heading('手速测试器周报', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 添加日期范围
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date_para.add_run('2026-05-08 至 2026-05-15')
    date_run.font.size = Pt(14)
    date_run.font.color.rgb = RGBColor(100, 100, 100)

    doc.add_paragraph()

    # 1. 本周修改总览
    doc.add_heading('一、本周修改总览', level=1)

    overview_para = doc.add_paragraph()
    overview_para.add_run('本周共有 ').bold = False
    overview_para.add_run('1').bold = True
    overview_para.add_run(' 次提交，主要完成了手速测试器的综合优化工作。')

    summary_items = [
        'Bug修复：修复了除零错误和NaN值处理问题',
        '性能优化：改进数据处理逻辑，过滤无效数据',
        '功能增强：添加最佳成绩记录和成就系统',
        'UX改进：添加键盘快捷键支持',
        'UI增强：添加测试中脉冲动画效果'
    ]

    for item in summary_items:
        p = doc.add_paragraph(item, style='List Bullet')

    doc.add_paragraph()

    # 2. 详细修改列表
    doc.add_heading('二、详细修改列表', level=1)

    commit_info = doc.add_paragraph()
    commit_info.add_run('提交信息：').bold = True
    commit_info.add_run('feat: 优化手速测试器 - Bug修复、性能改进和新功能增强\n')
    commit_info.add_run('提交者：').bold = True
    commit_info.add_run('AI Assistant\n')
    commit_info.add_run('提交日期：').bold = True
    commit_info.add_run('2026-05-13\n')
    commit_info.add_run('提交Hash：').bold = True
    commit_info.add_run('3f0de45')

    doc.add_paragraph()

    # 3. Bug修复详情
    doc.add_heading('三、Bug修复详情', level=1)

    bug_fixes = [
        ('除零错误修复', '修复了 exportImage 函数中的除零错误问题，防止在计算过程中出现除数为零的情况'),
        ('数组空值处理', '修复了 showTestDetails 函数中的数组空值处理问题，避免数组操作时出现异常'),
        ('NaN值过滤', '修复了 avgSpeeds 计算中的 NaN 值过滤问题，确保平均速度计算的准确性')
    ]

    for i, (title, desc) in enumerate(bug_fixes, 1):
        p = doc.add_paragraph()
        p.add_run(f'{i}. {title}').bold = True
        doc.add_paragraph(f'   描述：{desc}')

    doc.add_paragraph()

    # 4. 功能改进说明
    doc.add_heading('四、功能改进说明', level=1)

    improvements = [
        ('性能优化', '改进数据处理逻辑，过滤无效数据，提升整体运行效率'),
        ('最佳成绩记录', '新增最佳成绩记录功能，用户可以查看历史最佳表现'),
        ('成就系统', '添加成就系统，激励用户提升手速'),
        ('键盘快捷键', '添加键盘快捷键支持：空格键开始测试、ESC键关闭弹窗、Ctrl+R重置'),
        ('UI动画增强', '添加测试中的脉冲动画效果，提升用户视觉体验')
    ]

    for title, desc in improvements:
        p = doc.add_paragraph()
        p.add_run(f'• {title}：').bold = True
        p.add_run(desc)

    doc.add_paragraph()

    # 5. 下周优化建议
    doc.add_heading('五、下周优化建议', level=1)

    suggestions = [
        '继续完善成就系统，添加更多有趣的成就徽章',
        '优化移动端用户体验，确保在各种设备上流畅运行',
        '添加数据导出功能，支持用户导出测试历史记录',
        '考虑添加多人竞技模式，增加社交互动',
        '持续优化性能，减少页面加载时间',
        '收集用户反馈，针对性改进核心功能'
    ]

    for suggestion in suggestions:
        doc.add_paragraph(suggestion, style='List Bullet')

    doc.add_paragraph()

    # 页脚信息
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_run = footer.add_run(f'报告生成时间：{datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
    footer_run.font.size = Pt(9)
    footer_run.font.color.rgb = RGBColor(150, 150, 150)

    return doc

if __name__ == '__main__':
    # 创建周报
    doc = create_weekly_report()

    # 保存文件
    output_path = '/workspace/AI操作/documents/周报/手速测试器周报_20260508_20260515.docx'
    doc.save(output_path)

    print(f'周报已成功生成并保存至: {output_path}')
